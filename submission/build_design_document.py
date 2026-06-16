from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(__file__).resolve().parent / "03_设计文档" / "VESC2026_IMU手势识别设计文档.docx"
FLOW_IMAGE = ROOT / "VeriHealthi_QEMU_SDK_v3.7" / "流程图.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
LIGHT_FILL = "F2F4F7"
FONT = "Microsoft YaHei"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header=True):
    if header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if header and row_index == 0:
                set_cell_fill(cell, LIGHT_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=header and row_index == 0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("VESC 2026 | VeriHealthi IMU Gesture Recognition")
    set_run_font(header_run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("初赛设计文档 | 2026-06-15")
    set_run_font(footer_run, size=9, color=MUTED)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("VESC 2026 IMU 手势识别系统")
    set_run_font(run, size=23, bold=True, color=RGBColor(0, 0, 0))
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("基于 VeriHealthi SDK、随机森林与事件状态机的初赛设计文档")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("赛区", "【提交前填写】"),
        ("学校", "【提交前填写】"),
        ("队伍", "【提交前填写】"),
        ("队长", "【提交前填写】"),
        ("版本", "Final Candidate 1.0 | 2026-06-15"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}：")
        set_run_font(r, bold=True)
        r = p.add_run(value)
        set_run_font(r)

    doc.add_page_break()

    doc.add_heading("1. 项目概述", level=1)
    add_body(doc, "本项目在 VeriHealthi QEMU SDK 上完成 50 Hz 六轴 IMU 数据采集、任务间 Event 传输、CRC-8/SMBus 校验和四类有效手势识别。有效手势包括双指互点、握拳、抬腕和放下；其它动作不打印结果。")
    add_body(doc, "算法采用 1 秒滑动窗口、123 个整数时域特征和 50 棵随机森林。板端加入分类别置信阈值、连续窗口确认、全局冷却和抬腕状态机，减少误报与重复输出。")

    doc.add_heading("2. 赛题要求完成情况", level=1)
    add_table(
        doc,
        ["要求", "实现位置", "状态"],
        [
            ("创建 imu_task 与 algo_task", "main.c、imu_pipeline.c、algo_task.c", "完成"),
            ("IMU HAL 初始化并配置 50 Hz", "imu_pipeline.c", "完成"),
            ("中断触发数据读取并暂存 RAM", "imu_pipeline.c", "完成"),
            ("通过 Event 传输窗口", "EVENT_SEN_DATA_READY", "完成"),
            ("前 320000 Byte CRC-8/SMBus", "imu_pipeline.c", "完成"),
            ("algo manager 处理 Event", "algo_task.c", "完成"),
            ("按“时间ms, 手势”打印", "algo_task.c", "完成"),
            ("其它动作不打印", "后处理状态机", "完成"),
        ],
        [3300, 4300, 1760],
    )

    doc.add_heading("3. 系统架构", level=1)
    add_body(doc, "系统由初始化入口、IMU 采集任务、算法任务和随机森林推理模块组成。中断回调只更新 data-ready 计数，不调用 HAL、OS、CRC 或 printf，满足 ISR 使用约束。")
    if FLOW_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(FLOW_IMAGE), width=Inches(6.2))
        picture._inline.docPr.set("title", "系统软件流程图")
        picture._inline.docPr.set("descr", "IMU 中断、imu_task、Event、algo_task 和随机森林推理的数据流")
        caption = doc.add_paragraph("图 1  双任务与 Event 数据流")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in caption.runs:
            set_run_font(r, size=9.5, color=MUTED)

    doc.add_heading("3.1 IMU 采集链路", level=2)
    for text in [
        "通过 hal_imu_get_device 获取 IMU 实例并完成电源、默认参数、量程、ODR、FIFO 和中断配置。",
        "采样率设置为 50 Hz，加速度量程为 ±8G，陀螺仪量程为 ±2000 dps。",
        "每次读取一个 ImuGyroAccelData，写入 50 点环形缓冲区。",
        "窗口长度为 50 点，步长为 10 点，即每 200 ms 产生一个 1 秒窗口。",
        "窗口快照通过 EVENT_SEN_DATA_READY 发送给 algo_task。",
    ]:
        add_number(doc, text)

    doc.add_heading("3.2 CRC 校验", level=2)
    add_body(doc, "前 320000 Byte 数据按 ImuGyroAccelData 完整结构体字节序列输入 CRC-8/SMBus。参数为 poly=0x07、init=0x00、refin=false、refout=false、xorout=0x00。达到目标字节数后打印一次校验值。")

    doc.add_heading("4. 算法设计", level=1)
    doc.add_heading("4.1 数据整理与无泄漏划分", level=2)
    add_body(doc, "数据集共有 6 名用户。文件名包含 updown 的录制在 up 和 down 目录中重复出现，程序只读取一份原始 IMU 数据并合并两类标签。评估采用按用户留一交叉验证，训练集和验证集的用户完全隔离。")

    doc.add_heading("4.2 特征设计", level=2)
    add_table(
        doc,
        ["特征组", "数量", "作用"],
        [
            ("六轴范围、加速度均值、首尾差", "15", "描述整体幅度、姿态和方向变化"),
            ("动态绝对均值、范围和", "4", "描述动作总体强度"),
            ("平均绝对偏差", "6", "降低个人静态偏置影响"),
            ("前后半段及首尾四分之一差", "12", "描述运动趋势和翻转方向"),
            ("去均值最大偏差", "6", "捕捉瞬时冲击"),
            ("五段均值与五段范围", "60", "保留粗粒度时间形状"),
            ("五段运动包络均值与峰值", "20", "区分瞬时互点与持续握拳"),
        ],
        [3400, 900, 5060],
    )
    add_body(doc, "全部特征使用 int32 整数计算。Python 和 C 使用相同的对称四舍五入规则，避免导出后边界漂移。")

    doc.add_heading("4.3 随机森林", level=2)
    add_table(
        doc,
        ["参数", "取值"],
        [
            ("树数量", "50"),
            ("最大深度", "12"),
            ("最小叶节点样本数", "4"),
            ("类别权重", "balanced_subsample"),
            ("随机种子", "2026"),
            ("节点数", "76,106"),
        ],
        [4200, 5160],
    )

    doc.add_heading("4.4 输出状态机", level=2)
    add_bullet(doc, "投票阈值：pinch 42/50，clench 34/50，up/down 32/50。")
    add_bullet(doc, "同一类别连续 3 个窗口满足条件后才确认。")
    add_bullet(doc, "有效输出后进入 600 ms 全局冷却，并通过同类阻塞释放抑制单动作重复输出。")
    add_bullet(doc, "抬腕后才允许输出 down；输出 down 后恢复放下状态。")
    add_bullet(doc, "others 不产生打印。")

    doc.add_heading("5. 软件工作流程", level=1)
    for text in [
        "main 完成 SoC、Board 和 IMU pipeline 初始化。",
        "创建 imu_task 和 algo_task，调度器开始运行。",
        "algo_task 创建 manager 并注册 EVENT_SEN_DATA_READY。",
        "imu_task 等待算法任务就绪，随后持续读取 50 Hz IMU。",
        "每累计一个窗口步长，复制 50 点窗口并发送 Event。",
        "algo_task 提取特征、执行 50 棵树、运行状态机。",
        "识别有效手势后打印累计处理时间和类别。",
    ]:
        add_number(doc, text)

    doc.add_heading("6. 验证方法与结果", level=1)
    add_body(doc, "事件级评估严格复现主办方窗口规则。每个标签窗口从标签前 0.3 秒开始，到下一标签前 0.3 秒结束；窗口内多余输出均计为错误。")
    add_table(
        doc,
        ["类别", "TP", "FP", "FN", "精确率", "召回率", "F1"],
        [
            ("pinch", 2232, 354, 548, "86.31%", "80.29%", "83.19%"),
            ("clench", 2471, 176, 307, "93.35%", "88.95%", "91.10%"),
            ("up", 2326, 23, 65, "99.02%", "97.28%", "98.14%"),
            ("down", 2301, 39, 85, "98.33%", "96.44%", "97.38%"),
        ],
        [1300, 850, 850, 850, 1500, 1500, 1510],
    )
    add_body(doc, "六用户留一事件级宏 F1 为 92.45%。六折窗口准确率范围为 90.58% 至 93.52%。正确事件平均输出延迟约 318 ms。")
    add_body(doc, "公开数据仅用于开发与交叉验证，隐藏测试集成绩及最终名次以主办方测评为准。")

    doc.add_heading("7. 编译与资源", level=1)
    add_table(
        doc,
        ["项目", "结果"],
        [
            ("编译器", "Nuclei RISC-V GCC 14.2.1"),
            ("编译选项", "-O2 -Werror -Wall"),
            ("构建结果", "0 errors, 0 warnings"),
            ("text", "1,145,480 Byte"),
            ("data", "57,124 Byte"),
            ("bss", "30,128 Byte"),
            ("Flash 容量", "3 MB"),
            ("RAM 容量", "256 KB"),
        ],
        [4100, 5260],
    )
    add_body(doc, "模型数组定义为 const，链接到 Flash。运行时主要 RAM 包括任务栈、50 点窗口、Event 缓冲和少量特征统计数组。")

    doc.add_heading("8. 代码与复现", level=1)
    add_table(
        doc,
        ["文件", "作用"],
        [
            ("imu_pipeline.c/.h", "IMU 初始化、读取、窗口、Event 和 CRC"),
            ("algo_task.c/.h", "算法 manager 与结果打印"),
            ("gesture_algo.c/.h", "123 特征与后处理状态机"),
            ("gesture_rf_model.c/.h", "自动生成的随机森林数组与推理"),
            ("optimize_gesture_model.py", "无泄漏实验与官方窗口评测"),
            ("train_final_model.py", "全量训练及 C 模型导出"),
            ("verify_exported_model.py", "Python/C 森林票数一致性检查"),
        ],
        [3600, 5760],
    )
    add_body(doc, "复现顺序：安装 numpy、scikit-learn、joblib；运行 train_final_model.py；运行 verify_exported_model.py；在 NucleiStudio 中 Clean Project、Build Project，再使用 debug_qemu 调试。")

    doc.add_heading("9. 已知风险与后续优化", level=1)
    add_bullet(doc, "pinch 与 clench 的跨用户动作差异仍是主要误差来源。")
    add_bullet(doc, "后续可补充更多用户和佩戴松紧条件的数据，优先做困难样本挖掘。")
    add_bullet(doc, "可在不改变板端接口的前提下重新训练并导出模型。")
    add_bullet(doc, "正式提交前需完整跑完 QEMU 数据，核对 CRC 和输出格式。")

    doc.add_heading("10. 团队介绍与分工（待填写）", level=1)
    add_body(doc, "赛区：【填写】    学校：【填写】    队伍：【填写】    队长：【填写】")
    add_table(
        doc,
        ["姓名", "学号", "主要分工", "工作占比"],
        [("", "", "", ""), ("", "", "", ""), ("", "", "", "")],
        [1900, 1900, 3760, 1800],
    )
    add_body(doc, "指导教师：【填写】    联系方式：【填写】")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
